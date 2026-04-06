import os
from docx import Document

DOCX_PATH = os.path.join(os.getcwd(), "NLP.docx")

WELCOME_CANDIDATES = [
    os.path.join(os.getcwd(), "Welcome.txt"),
    os.path.join(os.path.expanduser("~"), "Desktop", "Text Mining and NLP", "Welcome.txt"),
]

OUTPUT_PATH = os.path.join(os.getcwd(), "NLP_updated.docx")

def read_welcome_text():
    for p in WELCOME_CANDIDATES:
        if os.path.isfile(p):
            with open(p, "r", encoding="utf-8") as f:
                return f.read().strip()
    raise FileNotFoundError(
        "Welcome.txt not found. Checked:\n  - " + "\n  - ".join(WELCOME_CANDIDATES)
    )
def main():
    # Open the docx
    doc = Document(DOCX_PATH)

    # 1) How many paragraphs?
    total_paras = len(doc.paragraphs)
    print(f"Total paragraphs (before): {total_paras}")

    # 2) For each paragraph, print the number of words it contains
    for i, p in enumerate(doc.paragraphs, start=1):
        wc = len(p.text.split())
        print(f"Paragraph {i}: {wc} words")

         # 3) Read content from Welcome.txt
    welcome_content = read_welcome_text()

    if doc.paragraphs:
        doc.paragraphs[0].insert_paragraph_before(welcome_content)
    else:
        # Edge case: no paragraphs in the doc
        doc.add_paragraph(welcome_content)

        doc.save(OUTPUT_PATH)
    print(f"Updated document saved to: {OUTPUT_PATH}")

    if __name__ == "__main__":
     main()
